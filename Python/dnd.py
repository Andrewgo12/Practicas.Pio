from docx import Document

# Crear un nuevo documento de Word
doc = Document()

# Agregar título
doc.add_heading('CSS Selectors', 0)

# Información de Selectores Básicos
doc.add_heading('1. Selectores Básicos', level=1)

# Element Selector
doc.add_heading('Element Selector', level=2)
doc.add_paragraph('''
p {
    color: blue;
}
Descripción: Aplica estilos a todos los elementos <p>.
''')

# Class Selector
doc.add_heading('Class Selector', level=2)
doc.add_paragraph('''
.example {
    color: red;
}
Descripción: Aplica estilos a todos los elementos con la clase example.
''')

# ID Selector
doc.add_heading('ID Selector', level=2)
doc.add_paragraph('''
#unique {
    color: green;
}
Descripción: Aplica estilos al elemento con el ID unique. (Recuerda que los IDs deben ser únicos en una página.)
''')

# Universal Selector
doc.add_heading('Universal Selector', level=2)
doc.add_paragraph('''
* {
    box-sizing: border-box;
}
Descripción: Aplica estilos a todos los elementos en la página.
''')

# Group Selector
doc.add_heading('Group Selector', level=2)
doc.add_paragraph('''
h1, h2, h3 {
    font-family: 'Helvetica', sans-serif;
}
Descripción: Aplica el mismo estilo a varios tipos de elementos.
''')

# Class Multiple
doc.add_heading('Class Multiple', level=2)
doc.add_paragraph('''
.box, .circle {
    border: 2px solid black;
}
Descripción: Aplica estilos a elementos con varias clases.
''')

# ID and Class Combined
doc.add_heading('ID and Class Combined', level=2)
doc.add_paragraph('''
#header .title {
    color: #ff5733;
}
Descripción: Aplica estilos a elementos con una clase dentro de un ID específico.
''')

# Attribute with Specific Value
doc.add_heading('Attribute with Specific Value', level=2)
doc.add_paragraph('''
[type="email"] {
    background-color: #f0f8ff;
}
Descripción: Aplica estilos a campos de entrada de tipo email.
''')

# Información de Selectores Combinados
doc.add_heading('2. Selectores Combinados', level=1)

# Descendant Selector
doc.add_heading('Descendant Selector', level=2)
doc.add_paragraph('''
.container p {
    color: gray;
}
Descripción: Aplica estilos a todos los <p> dentro de un elemento con la clase container.
''')

# Child Selector
doc.add_heading('Child Selector', level=2)
doc.add_paragraph('''
.container > p {
    color: yellow;
}
Descripción: Aplica estilos solo a los <p> que son hijos directos de un elemento con la clase container.
''')

# Adjacent Sibling Selector
doc.add_heading('Adjacent Sibling Selector', level=2)
doc.add_paragraph('''
h2 + p {
    color: pink;
}
Descripción: Aplica estilos al primer <p> que sigue inmediatamente a un <h2>.
''')

# General Sibling Selector
doc.add_heading('General Sibling Selector', level=2)
doc.add_paragraph('''
h2 ~ p {
    color: orange;
}
Descripción: Aplica estilos a todos los <p> que siguen a un <h2>, no necesariamente de inmediato.
''')

# Repetir la información (doblar)
# Seleccionaremos solo parte del contenido, repetiremos algunos ejemplos de pseudoclases y pseudoelementos
# para mostrar cómo duplicar

# Pseudoclases: :hover
doc.add_heading('4. Pseudoclases', level=1)
doc.add_heading('a:hover', level=2)
doc.add_paragraph('''
a:hover {
    color: red;
}
Descripción: Cambia el color de los enlaces a rojo cuando se pasa el ratón sobre ellos.
''')

# Duplicado de pseudoclases
doc.add_heading('a:hover (duplicado)', level=2)
doc.add_paragraph('''
a:hover {
    color: red;
}
Descripción: Cambia el color de los enlaces a rojo cuando se pasa el ratón sobre ellos.
''')

# Pseudoelementos: ::before
doc.add_heading('5. Pseudoelementos', level=1)
doc.add_heading('p::before', level=2)
doc.add_paragraph('''
p::before {
    content: "★ ";
    color: gold;
}
Descripción: Añade una estrella dorada antes del contenido de cada párrafo.
''')

# Duplicado de pseudoelementos
doc.add_heading('p::before (duplicado)', level=2)
doc.add_paragraph('''
p::before {
    content: "★ ";
    color: gold;
}
Descripción: Añade una estrella dorada antes del contenido de cada párrafo.
''')

# Guardar el documento
file_path = "/mnt/data/CSS_selectores_doblados.docx"
doc.save(file_path)

file_path