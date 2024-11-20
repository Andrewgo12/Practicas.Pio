from docx import Document

# Create a new Document
doc = Document()
doc.add_heading('Ejemplos de Integrales Dobles y Triples', 0)

# Add heading for double integrals examples
doc.add_heading('Ejemplos de Integrales Dobles', level=1)

# List of double integral examples
double_integrals = [
    "1. Función: z = x^2 + y^2\n   Región delimitada por las siguientes curvas: y = x, y = 1 - x",
    "2. Función: z = x^2 + y^2\n   Región delimitada por el círculo: x^2 + y^2 = 4",
    "3. Función: z = 3x + 2y\n   Región delimitada por las siguientes rectas: y = x, y = 2 + x",
    "4. Función: z = x^2 + y^2\n   Región delimitada por el triángulo: y = 0, y = 3x, x = 1",
    "5. Función: z = 2x + y^2\n   Región delimitada por las siguientes curvas: y = 0, y = 4, x = 0, x = 3",
    "6. Función: z = x + y\n   Región delimitada por el círculo: x^2 + y^2 = 1",
    "7. Función: z = 2x + y^2\n   Región delimitada por las siguientes rectas: y = 2x, y = x + 1",
    "8. Función: z = x^2 + y^2\n   Región delimitada por el triángulo: y = x, y = 2 - x, x = 1",
    "9. Función: z = 5x + 3y\n   Región delimitada por las siguientes rectas: y = x, y = 1 - x",
    "10. Función: z = x^2 + y^2\n   Región delimitada por la elipse: x^2/4 + y^2/9 = 1"
]

# Add examples to the document
for example in double_integrals:
    doc.add_paragraph(example)

# Add heading for triple integrals examples
doc.add_heading('Ejemplos de Integrales Triples', level=1)

# List of triple integral examples
triple_integrals = [
    "1. Función: f(x, y, z) = x^2 + y^2 + z^2\n   Región delimitada por: x^2 + y^2 + z^2 ≤ 4",
    "2. Función: f(x, y, z) = e^(-x^2 - y^2 - z^2)\n   Región delimitada por: x^2 + y^2 + z^2 ≤ 1",
    "3. Función: f(x, y, z) = x + y + z\n   Región delimitada por un cubo: 0 ≤ x ≤ 2, 0 ≤ y ≤ 2, 0 ≤ z ≤ 2",
    "4. Función: f(x, y, z) = 3x + 2y + z^2\n   Región delimitada por: x^2 + y^2 ≤ 4, z ∈ [0, 2]",
    "5. Función: f(x, y, z) = 2x + 3y + 4z\n   Región delimitada por: 0 ≤ x ≤ 1, 0 ≤ y ≤ 2, 0 ≤ z ≤ 3",
    "6. Función: f(x, y, z) = x^3 + y^3 + z^3\n   Región delimitada por: x^2 + y^2 + z^2 ≤ 1",
    "7. Función: f(x, y, z) = sin(x) + cos(y) + z\n   Región delimitada por: x ∈ [0, π], y ∈ [0, 2π], z ∈ [0, 1]",
    "8. Función: f(x, y, z) = x^2 + y^2 + z^2\n   Región delimitada por un cilindro: x^2 + y^2 ≤ 4, z ∈ [0, 2]",
    "9. Función: f(x, y, z) = e^(-x^2 - y^2)\n   Región delimitada por: -1 ≤ x ≤ 1, -1 ≤ y ≤ 1, 0 ≤ z ≤ 1",
    "10. Función: f(x, y, z) = 5x + 2y + 3z\n   Región delimitada por: x ∈ [0, 1], y ∈ [0, 1], z ∈ [0, 1]"
]

# Add examples to the document
for example in triple_integrals:
    doc.add_paragraph(example)

# Save the document
file_path = "/mnt/data/Ejemplos_Integrales_Dobles_y_Triples.docx"
doc.save(file_path)

file_path
