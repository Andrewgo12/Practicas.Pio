from docx import Document

# Crear un nuevo documento
doc = Document()
doc.add_heading('Cálculos de Campo Eléctrico, Fuerza y Energía Potencial', level=1)

# Cálculo del campo eléctrico
doc.add_heading('a) Campo Eléctrico en el Centro de la Circunferencia', level=2)
doc.add_paragraph('Datos:')
doc.add_paragraph('- q1 = 3 µC')
doc.add_paragraph('- q2 = -7 µC')
doc.add_paragraph('- q3 = 4 µC')
doc.add_paragraph('- R = 8 cm = 0.08 m')

doc.add_paragraph('Campo eléctrico de cada carga:')

doc.add_paragraph('1. Campo eléctrico debido a q1:')
doc.add_paragraph('E1 = k * q1 / R^2')
doc.add_paragraph('E1 = (8.99 * 10^9 N·m^2/C^2 * 3 * 10^-6 C) / (0.08 m)^2 = 4.22 * 10^6 N/C')
doc.add_paragraph('Componentes:')
doc.add_paragraph('E1x = 4.22 * 10^6 N/C')
doc.add_paragraph('E1y = 0')

doc.add_paragraph('2. Campo eléctrico debido a q2:')
doc.add_paragraph('Ángulo: 30° con respecto al eje -x.')
doc.add_paragraph('E2 = k * |q2| / R^2')
doc.add_paragraph('E2 = (8.99 * 10^9 N·m^2/C^2 * 7 * 10^-6 C) / (0.08 m)^2 = 9.85 * 10^6 N/C')
doc.add_paragraph('Componentes:')
doc.add_paragraph('E2x = -9.85 * 10^6 * cos(30°) = -8.54 * 10^6 N/C')
doc.add_paragraph('E2y = -9.85 * 10^6 * sin(30°) = -4.93 * 10^6 N/C')

doc.add_paragraph('3. Campo eléctrico debido a q3:')
doc.add_paragraph('Ángulo: 50° con respecto al eje -y.')
doc.add_paragraph('E3 = k * q3 / R^2')
doc.add_paragraph('E3 = (8.99 * 10^9 N·m^2/C^2 * 4 * 10^-6 C) / (0.08 m)^2 = 1.12 * 10^7 N/C')
doc.add_paragraph('Componentes:')
doc.add_paragraph('E3x = 1.12 * 10^7 * sin(50°) = 8.57 * 10^6 N/C')
doc.add_paragraph('E3y = -1.12 * 10^7 * cos(50°) = -7.15 * 10^6 N/C')

doc.add_paragraph('Campo eléctrico total en el centro:')
doc.add_paragraph('Ex = E1x + E2x + E3x')
doc.add_paragraph('Ex = 4.22 * 10^6 - 8.54 * 10^6 + 8.57 * 10^6 = 4.25 * 10^6 N/C')
doc.add_paragraph('Ey = E1y + E2y + E3y')
doc.add_paragraph('Ey = 0 - 4.93 * 10^6 - 7.15 * 10^6 = -1.21 * 10^7 N/C')

# Cálculo de la fuerza
doc.add_heading('b) Fuerza Eléctrica sobre Q = -10 µC en el Centro', level=2)
doc.add_paragraph('Fx = Q * Ex')
doc.add_paragraph('Fx = -10 * 10^-6 C * 4.25 * 10^6 N/C = -42.5 N')
doc.add_paragraph('Fy = Q * Ey')
doc.add_paragraph('Fy = -10 * 10^-6 C * (-1.21 * 10^7 N/C) = 121.0 N')

doc.add_paragraph('Fuerza total:')
doc.add_paragraph('F = sqrt(Fx^2 + Fy^2)')
doc.add_paragraph('F = sqrt((-42.5)^2 + 121.0^2) = 128.5 N')

doc.add_paragraph('Ángulo:')
doc.add_paragraph('θ = tan^-1(Fy / Fx)')
doc.add_paragraph('θ = tan^-1(121.0 / -42.5) = 71.4° (desde el eje -x)')

# Cálculo de la energía potencial
doc.add_heading('c) Energía Potencial Eléctrica de la Carga Q', level=2)
doc.add_paragraph('U = k * Q * (q1/R + q2/R + q3/R)')
doc.add_paragraph('U = 8.99 * 10^9 N·m^2/C^2 * (-10 * 10^-6 C) * ( (3 * 10^-6 + (-7 * 10^-6) + 4 * 10^-6) / 0.08 m )')
doc.add_paragraph('U = 0')

# Guardar el documento
doc.save('calculos_electricos.docx')

print("Documento Word creado exitosamente.")